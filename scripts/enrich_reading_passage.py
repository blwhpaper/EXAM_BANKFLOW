#!/usr/bin/env python3
"""
enrich_reading_passage.py
为 section_type=reading 的记录提取并写入 passage_text 字段。

规则：
1. 只处理 reading 记录
2. 按 exam_id + source_file 分组
3. 在原始 docx 中定位 A/B/C/D 篇章标签
4. 每篇文章范围 = 标签下一段 到 该篇第一道题前一段
5. 跳过并标记 dirty：
   - options 少于 4 个
   - 已有 data_quality="dirty"
   - passage_text_verified=True
"""

from __future__ import annotations

import argparse
import json
import re
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
JSONL_PATH = PROJECT_ROOT / "workflow/records/EXAM-CLEAN-010_STRUCTURED_QUESTION_BANK.jsonl"

LABEL_RE = re.compile(r"^[A-D]\.?$")
QUESTION_RE = re.compile(r"^\d+[\.\s]")
PARAGRAPH_NUM_RE = re.compile(r"\d+")
INSTRUCTION_PATTERNS = (
    "阅读下列短文",
    "阅读下面短文",
    "阅读理解",
    "第一节",
    "第二节",
    "从每题所给的",
    "从短文后的选项中",
    "选出最佳选项",
    "回答下列问题",
)


@dataclass
class PassageSlice:
    label: str
    start_idx: int
    end_idx: int
    paragraphs: list[str]


@dataclass
class QuestionBlock:
    start_idx: int
    end_idx: int


def load_records(jsonl_path: Path) -> list[dict]:
    records: list[dict] = []
    with jsonl_path.open(encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                records.append(json.loads(line))
    return records


def write_records(jsonl_path: Path, records: list[dict]) -> None:
    tmp_path = jsonl_path.with_suffix(".tmp")
    with tmp_path.open("w", encoding="utf-8") as f:
        for record in records:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
    tmp_path.replace(jsonl_path)


def normalize_paragraph(text: str) -> str:
    return " ".join(text.replace("\t", " ").split())


def is_instruction_line(text: str) -> bool:
    return any(pattern in text for pattern in INSTRUCTION_PATTERNS)


def build_question_blocks(question_indices: list[int], total_paragraphs: int) -> list[QuestionBlock]:
    blocks: list[QuestionBlock] = []
    if not question_indices:
        return blocks

    block_start = question_indices[0]
    prev = question_indices[0]
    for idx in question_indices[1:]:
        if idx - prev > 4:
            blocks.append(QuestionBlock(start_idx=block_start, end_idx=prev))
            block_start = idx
        prev = idx
    blocks.append(QuestionBlock(start_idx=block_start, end_idx=prev))
    return blocks


def load_docx_paragraphs(docx_path: Path) -> list[str]:
    return [normalize_paragraph(p.text) for p in Document(str(docx_path)).paragraphs]


def extract_passage_slices(paragraphs: list[str], first_reading_question_idx: int | None = None) -> list[PassageSlice]:
    label_indices = [(idx, text.rstrip(".")) for idx, text in enumerate(paragraphs) if LABEL_RE.match(text)]
    question_indices = [idx for idx, text in enumerate(paragraphs) if QUESTION_RE.match(text)]

    slices: list[PassageSlice] = []

    if first_reading_question_idx is not None:
        prior_labels = [idx for idx, _ in label_indices if idx < first_reading_question_idx]
        if not prior_labels:
            implicit_start = None
            for idx in range(first_reading_question_idx - 1, -1, -1):
                text = paragraphs[idx].strip()
                if "阅读下列短文" in text or "阅读下面短文" in text:
                    implicit_start = idx + 1
                    break
            if implicit_start is not None and implicit_start < first_reading_question_idx:
                kept = []
                for idx in range(implicit_start, first_reading_question_idx):
                    text = paragraphs[idx].strip()
                    if not text:
                        continue
                    if is_instruction_line(text):
                        continue
                    kept.append(text)
                if kept:
                    slices.append(
                        PassageSlice(
                            label="A",
                            start_idx=implicit_start,
                            end_idx=first_reading_question_idx - 1,
                            paragraphs=kept,
                        )
                    )

    for label_idx, label in label_indices:
        next_questions = [idx for idx in question_indices if idx > label_idx]
        if not next_questions:
            continue

        first_question_idx = next_questions[0]
        start_idx = label_idx + 1
        end_idx = first_question_idx - 1
        if end_idx < start_idx:
            continue

        kept = []
        for idx in range(start_idx, end_idx + 1):
            text = paragraphs[idx].strip()
            if not text:
                continue
            if is_instruction_line(text):
                continue
            kept.append(text)

        slices.append(
            PassageSlice(
                label=label,
                start_idx=start_idx,
                end_idx=end_idx,
                paragraphs=kept,
            )
        )
    return slices


def parse_paragraph_index(record: dict) -> int:
    paragraph = str(record["source_span"]["paragraph"])
    match = PARAGRAPH_NUM_RE.search(paragraph)
    if not match:
        raise ValueError(f"invalid paragraph locator: {paragraph!r}")
    return int(match.group())


def count_options(record: dict) -> int:
    options = record.get("options") or []
    return len(options) if isinstance(options, list) else 0


def mark_dirty(record: dict) -> bool:
    if record.get("passage_text_verified") is True:
        return False
    if record.get("data_quality") == "dirty":
        return False
    record["data_quality"] = "dirty"
    return True


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--exam", help="只处理指定 exam_id")
    parser.add_argument("--force", action="store_true", help="强制覆盖已有 passage_text")
    parser.add_argument("--dry-run", action="store_true", help="只打印不写回")
    args = parser.parse_args()

    records = load_records(JSONL_PATH)
    id_to_record = {record["question_id"]: record for record in records}

    reading_records = [r for r in records if r.get("section_type") == "reading"]
    if args.exam:
        reading_records = [r for r in reading_records if r.get("exam_id") == args.exam]

    groups: dict[tuple[str, str], list[dict]] = defaultdict(list)
    for record in reading_records:
        groups[(record["exam_id"], record.get("source_file", ""))].append(record)

    success_groups = 0
    skip_reasons: dict[str, int] = defaultdict(int)
    failed_groups: list[tuple[str, str]] = []
    dirty_marked = 0
    updated_records = 0

    for (exam_id, source_file), group in sorted(groups.items()):
        docx_path = PROJECT_ROOT / source_file
        group_label = f"{exam_id} | {source_file}"

        if not docx_path.exists():
            failed_groups.append((str(docx_path), "文件不存在"))
            print(f"[失败] {group_label}: 文件不存在")
            continue

        try:
            paragraphs = load_docx_paragraphs(docx_path)
        except Exception as exc:
            failed_groups.append((str(docx_path), f"docx 解析失败: {exc}"))
            print(f"[失败] {group_label}: docx 解析失败: {exc}")
            continue

        valid_candidates = []
        for record in group:
            if record.get("data_quality") == "dirty":
                skip_reasons["已有 data_quality=dirty"] += 1
                continue
            if record.get("passage_text_verified") is True:
                skip_reasons["passage_text_verified=True"] += 1
                continue
            if count_options(record) < 4:
                if mark_dirty(id_to_record[record["question_id"]]):
                    dirty_marked += 1
                skip_reasons["options 少于 4 个"] += 1
                continue
            valid_candidates.append(record)

        valid_candidates.sort(key=parse_paragraph_index)
        if not valid_candidates:
            print(f"[跳过] {group_label}: 无可处理 reading 记录")
            continue

        first_reading_question_idx = min(parse_paragraph_index(record) for record in valid_candidates)
        slices = extract_passage_slices(paragraphs, first_reading_question_idx=first_reading_question_idx)
        if not slices:
            failed_groups.append((str(docx_path), "未找到可用篇章标签"))
            print(f"[失败] {group_label}: 未找到可用篇章标签")
            continue

        first_question_indices = [parse_paragraph_index(record) for record in valid_candidates]
        label_to_slice: dict[str, PassageSlice] = {}
        for slice_info in slices:
            next_qs = [idx for idx in first_question_indices if idx > slice_info.end_idx]
            first_q_idx = next_qs[0] if next_qs else None
            if first_q_idx is None:
                continue
            label_to_slice[slice_info.label] = slice_info

        assigned = 0
        slice_order = sorted(label_to_slice.values(), key=lambda item: item.end_idx)
        for record in valid_candidates:
            paragraph_idx = parse_paragraph_index(record)
            matched = None
            for slice_info in slice_order:
                next_question_start = slice_info.end_idx + 1
                if paragraph_idx >= next_question_start:
                    matched = slice_info
            if matched is None:
                skip_reasons["未匹配到篇章"] += 1
                continue

            target = id_to_record[record["question_id"]]
            if not args.force and target.get("passage_text"):
                skip_reasons["已有 passage_text 且未使用 --force"] += 1
                continue

            passage_text = "\n".join(matched.paragraphs)
            if not passage_text:
                skip_reasons["提取结果为空"] += 1
                continue

            if not args.dry_run:
                target["passage_text"] = passage_text
            updated_records += 1
            assigned += 1

        if assigned == 0:
            print(f"[跳过] {group_label}: 无记录完成回填")
            continue

        success_groups += 1
        print(f"\n=== {exam_id} | {Path(source_file).name} ===")
        for slice_info in slice_order:
            preview = slice_info.paragraphs[:3]
            print(f"{slice_info.label}")
            for line in preview:
                print(line)
            print(f"总段落数: {len(slice_info.paragraphs)}")
            print()

    print("统计")
    print(f"成功 {success_groups} 组")
    print(f"跳过 {sum(skip_reasons.values())} 条")
    if skip_reasons:
        print("跳过原因：")
        for reason, count in sorted(skip_reasons.items()):
            print(f"{reason}: {count}")
    print(f"失败 {len(failed_groups)} 组")
    if failed_groups:
        print("失败文件：")
        for path, reason in failed_groups:
            print(f"{path} | {reason}")
    if dirty_marked:
        print(f"本次新增 dirty 标记: {dirty_marked} 条")

    if args.dry_run:
        print("[dry-run] 未写回 JSONL")
        return

    if updated_records or dirty_marked:
        write_records(JSONL_PATH, records)
        print(f"已写回: {JSONL_PATH}")
    else:
        print("无变更，不写回文件")


if __name__ == "__main__":
    main()
