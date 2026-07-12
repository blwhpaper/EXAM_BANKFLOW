#!/usr/bin/env python3
# =====================================
# !! FORMAT LOCK — 格式基准版本 !!
# 本脚本是唯一授权的Word导出实现
# 任何格式修改必须改此文件并同步specs/export_reference.py
# 禁止agent在调用时绕过此脚本自行生成docx
# 上次验证：E005 完形填空，4项格式全部通过
# =====================================
"""
scripts/export.py — 统一导出入口
用法：
  python scripts/export.py --type cloze --exam E005 --out /tmp/out.docx
  python scripts/export.py --type cloze --out /tmp/out.docx
  python scripts/export.py --type reading --exam E005 --out /tmp/out.docx
  python scripts/export.py --type reading_7to5 --exam E005 --out /tmp/out.docx
  python scripts/export.py --type cloze --subtype cloze_lexical --out /tmp/out.docx
  python scripts/export.py --type cloze --strategy context_clue --out /tmp/out.docx
"""

import argparse
import json
import re
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Twips
from docx.enum.section import WD_ORIENT

# ── 硬编码格式常量 ────────────────────────────────────────────────────────────
FONT_NAME   = "Times New Roman"
FONT_SIZE   = Pt(12)
PAGE_W      = 11906   # A4 宽（DXA）
PAGE_H      = 16838   # A4 高（DXA）
MARGIN      = 1440    # 四边边距（DXA）
LINE_RULE   = "auto"
LINE_VAL    = 360     # 1.5倍行距
SP_BEFORE   = 0
SP_AFTER    = 0
SP_AFTER_Q  = 200     # 答案行后间距
SP_AFTER_PASSAGE = 240  # 文章与题目块之间

TAB_STOPS   = [2200, 4400, 6600]  # 选项列位置（DXA），B/C/D 跳转点

JSONL_PATH  = Path("/Volumes/KIOXIA_1TB/01_PROJECTS/EXAM_BANKFLOW/"
                   "workflow/records/EXAM-CLEAN-010_STRUCTURED_QUESTION_BANK.jsonl")
QNUM_RE = re.compile(r"\d+")


# ── 工具函数 ──────────────────────────────────────────────────────────────────

def set_spacing(para, before=SP_BEFORE, after=SP_AFTER, line=LINE_VAL):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"),   str(before))
    spacing.set(qn("w:after"),    str(after))
    spacing.set(qn("w:line"),     str(line))
    spacing.set(qn("w:lineRule"), LINE_RULE)
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)


def add_tab_stops(para):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    for pos in TAB_STOPS:
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), str(pos))
        tabs.append(tab)
    existing = pPr.find(qn("w:tabs"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(tabs)


def make_run(para, text, bold=False):
    run = para.add_run(text)
    run.font.name       = FONT_NAME
    run.font.size       = FONT_SIZE
    run.font.bold       = bold
    # 中文字体
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    FONT_NAME)
    rFonts.set(qn("w:hAnsi"),    FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    return run


def add_passage_para(doc, text, after=SP_AFTER):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, before=SP_BEFORE, after=after)
    make_run(para, text)
    return para


def add_stem_para(doc, qnum, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, before=120, after=0)
    make_run(para, f"{qnum}. {text}")
    return para


def add_options_para(doc, options, has_answer, qnum=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_tab_stops(para)
    after = 0 if has_answer else SP_AFTER_Q
    set_spacing(para, before=120, after=after)

    labels = ["A", "B", "C", "D"]
    opt_map = {o["label"]: o["text"] for o in options}
    prefix = f"{qnum}. " if qnum is not None else ""
    text = prefix + "\t".join(f"{lbl}. {opt_map.get(lbl, '')}" for lbl in labels)
    make_run(para, text)
    return para


def add_answer_para(doc, answer):
    ans_text = answer.strip() if (answer and answer.strip()) else "待补充"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(para, before=0, after=SP_AFTER_Q)
    make_run(para, f"【答案】{ans_text}", bold=True)
    return para


def add_answer_block(doc, records):
    """整篇模式：在文章单元末尾输出答案汇总块，每5题一行。"""
    sorted_recs = sorted(records, key=lambda r: parse_qnum(r.get("source_question_number")))

    # 空行
    spacer = doc.add_paragraph()
    set_spacing(spacer, before=0, after=0)

    # 标题行
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(title, before=0, after=0)
    make_run(title, "【答案】", bold=True)

    # 每5题一行
    for i in range(0, len(sorted_recs), 5):
        chunk = sorted_recs[i:i + 5]
        nums = [parse_qnum(r.get("source_question_number")) for r in chunk]
        answers = [
            (r.get("answer") or "").strip() or "_"
            for r in chunk
        ]
        label = str(nums[0]) if len(nums) == 1 else f"{nums[0]}-{nums[-1]}"
        line_text = f"{label}  " + " ".join(answers)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(para, before=0, after=0)
        make_run(para, line_text, bold=True)


def add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)
    set_spacing(para, before=0, after=0)


# ── 文档初始化 ────────────────────────────────────────────────────────────────

def make_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Twips(PAGE_W)
    sec.page_height   = Twips(PAGE_H)
    sec.top_margin    = Twips(MARGIN)
    sec.bottom_margin = Twips(MARGIN)
    sec.left_margin   = Twips(MARGIN)
    sec.right_margin  = Twips(MARGIN)
    # 清除默认样式间距
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.space_before = 0
    style.paragraph_format.space_after  = 0
    return doc


# ── 主逻辑 ────────────────────────────────────────────────────────────────────

def load_records(path: Path) -> list:
    with open(path, encoding="utf-8") as f:
        return [json.loads(l) for l in f if l.strip()]


def parse_qnum(value) -> int:
    match = QNUM_RE.search(str(value or ""))
    return int(match.group()) if match else 0


def build_cloze_doc(groups: dict, subtype: str | None, strategy: str | None) -> Document:
    """
    groups: { (exam_id, source_file): [records sorted by q_number] }
    """
    doc        = make_document()
    first_unit = True

    for (exam_id, source_file), records in sorted(groups.items()):
        # 细分过滤
        if subtype:
            records = [r for r in records if r.get("question_subtype") == subtype]
        if strategy:
            records = [r for r in records if r.get("strategy") == strategy]
        if not records:
            continue

        if not first_unit:
            add_page_break(doc)
        first_unit = False

        use_excerpt = bool(subtype or strategy)

        if use_excerpt:
            # 细分模式：excerpt 上下文 + 选项行（含序号），末尾汇总答案
            for rec in records:
                excerpt = rec.get("source_span", {}).get("excerpt", rec.get("question_text", ""))
                add_passage_para(doc, excerpt, after=SP_AFTER_PASSAGE)
                add_options_para(doc, rec.get("options") or [], has_answer=False,
                                 qnum=rec["source_question_number"])
            add_answer_block(doc, records)
        else:
            # 整篇模式：文章原文 → 所有题目
            passage = records[0].get("passage_text", "")
            if passage:
                paras = [p for p in passage.split("\n") if p.strip()]
                for i, p in enumerate(paras):
                    after = SP_AFTER_PASSAGE if i == len(paras) - 1 else SP_AFTER
                    add_passage_para(doc, p, after=after)
            else:
                add_passage_para(doc, f"[{exam_id} 文章原文缺失]", after=SP_AFTER_PASSAGE)

            for rec in records:
                add_options_para(doc, rec.get("options") or [], has_answer=False,
                                 qnum=rec["source_question_number"])
            add_answer_block(doc, records)

    return doc


def add_7to5_options_block(doc, options):
    """七选五选项池：A-G 每项单独一行（句子较长，不分列）。"""
    spacer = doc.add_paragraph()
    set_spacing(spacer, before=0, after=0)
    for opt in options:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(para, before=0, after=0)
        make_run(para, f"{opt['label']}. {opt['text']}")


def split_passage_paragraphs(passage: str) -> list[str]:
    return [p.strip() for p in (passage or "").split("\n") if p.strip()]


def get_7to5_placeholder(rec: dict) -> str:
    question_text = rec.get("question_text") or ""
    match = re.search(r"____(\d+)____", question_text)
    if match:
        return f"____{match.group(1)}____"
    qnum = parse_qnum(rec.get("source_question_number"))
    return f"____{qnum}____" if qnum else ""


def find_7to5_para_index(paras: list[str], rec: dict) -> int:
    placeholder = get_7to5_placeholder(rec)
    if placeholder:
        for idx, para in enumerate(paras):
            if placeholder in para:
                return idx
    return max(len(paras) - 1, 0)


def add_7to5_marker_line(doc, records: list[dict]):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(para, before=0, after=SP_AFTER)
    chunks = []
    for rec in sorted(records, key=lambda r: parse_qnum(r.get("source_question_number"))):
        qnum = parse_qnum(rec.get("source_question_number"))
        methods = rec.get("solution_methods") or []
        label = "·".join(methods) if methods else "unlabeled"
        chunks.append(f"{qnum}【{label}】")
    make_run(para, "▶ " + "  ".join(chunks), bold=True)


def prepare_7to5_method_groups(groups: dict, method: str, limit: int | None):
    selected_ids: list[str] = []
    selected_set: set[str] = set()
    expanded_groups: dict[tuple, list] = {}

    for key in sorted(groups):
        group_records = groups[key]
        paras = split_passage_paragraphs(group_records[0].get("passage_text", ""))
        para_map: dict[int, list[dict]] = defaultdict(list)
        for rec in group_records:
            para_index = find_7to5_para_index(paras, rec)
            rec["_export_para_index"] = para_index
            para_map[para_index].append(rec)

        matched = [
            rec for rec in group_records
            if method in (rec.get("solution_methods") or [])
        ]
        matched.sort(key=lambda r: parse_qnum(r.get("source_question_number")))

        included_ids: set[str] = set()
        for rec in matched:
            rec_id = rec.get("question_id") or f"{key}-{rec.get('source_question_number')}"
            if limit is not None and len(selected_ids) >= limit:
                break
            selected_ids.append(rec_id)
            selected_set.add(rec_id)
            para_index = rec["_export_para_index"]
            for block_rec in para_map.get(para_index, []):
                block_id = block_rec.get("question_id") or f"{key}-{block_rec.get('source_question_number')}"
                included_ids.add(block_id)

        expanded = [
            rec for rec in group_records
            if (rec.get("question_id") or f"{key}-{rec.get('source_question_number')}") in included_ids
        ]
        expanded.sort(key=lambda r: (r["_export_para_index"], parse_qnum(r.get("source_question_number"))))
        if expanded:
            expanded_groups[key] = expanded

        if limit is not None and len(selected_ids) >= limit:
            # limit applies to method-matched questions; later groups are ignored
            break

    answer_records = []
    for key in sorted(expanded_groups):
        for rec in expanded_groups[key]:
            rec_id = rec.get("question_id") or f"{key}-{rec.get('source_question_number')}"
            if rec_id in selected_set:
                answer_records.append(rec)

    return expanded_groups, answer_records


def build_7to5_doc(groups: dict, method: str | None = None, answer_records: list[dict] | None = None) -> Document:
    """
    整篇模式（无 method 筛选）：
      passage_text → A-G 选项池 → 答案汇总块

    细分模式（有 method 筛选）：
      按段落块输出：从首段到空格所在段 → A-G 选项池 → 标注行；末尾统一答案汇总块
    """
    doc        = make_document()
    first_unit = True

    for (exam_id, source_file), records in sorted(groups.items()):
        if not records:
            continue
        if not first_unit:
            add_page_break(doc)
        first_unit = False

        opts = records[0].get("options") or []

        if method:
            paras = split_passage_paragraphs(records[0].get("passage_text", ""))
            blocks: dict[int, list[dict]] = defaultdict(list)
            for rec in records:
                para_index = rec.get("_export_para_index")
                if para_index is None:
                    para_index = find_7to5_para_index(paras, rec)
                blocks[para_index].append(rec)

            sorted_para_indices = sorted(blocks)
            # 文章中第一个含空白的段落（扫描全篇 passage，与导出集合无关）
            first_blank_para = next(
                (i for i, p in enumerate(paras) if re.search(r"____\d+____", p)),
                0
            )

            for block_idx, para_index in enumerate(sorted_para_indices):
                block_records = sorted(
                    blocks[para_index],
                    key=lambda r: parse_qnum(r.get("source_question_number"))
                )
                if paras:
                    # 该题所在段是文章第一个空白段：输出 paras[0..N]（需要前文上下文）
                    # 否则：只输出 paras[N]（该段即可）
                    if para_index == first_blank_para:
                        for idx in range(para_index + 1):
                            after = SP_AFTER if idx < para_index else 0
                            add_passage_para(doc, paras[idx], after=after)
                    else:
                        add_passage_para(doc, paras[para_index], after=0)
                else:
                    add_passage_para(doc, f"[{exam_id} 文章原文缺失]", after=0)

                spacer = doc.add_paragraph()
                set_spacing(spacer, before=0, after=0)
                if opts:
                    add_7to5_options_block(doc, opts)

                spacer = doc.add_paragraph()
                set_spacing(spacer, before=0, after=0)
                add_7to5_marker_line(doc, block_records)

                if block_idx != len(blocks) - 1:
                    spacer = doc.add_paragraph()
                    set_spacing(spacer, before=0, after=SP_AFTER_PASSAGE)
        else:
            passage = records[0].get("passage_text", "")
            if passage:
                paras = [p for p in passage.split("\n") if p.strip()]
                for i, p in enumerate(paras):
                    after = SP_AFTER_PASSAGE if i == len(paras) - 1 else SP_AFTER
                    add_passage_para(doc, p, after=after)
            else:
                add_passage_para(doc, f"[{exam_id} 文章原文缺失]", after=SP_AFTER_PASSAGE)

            if opts:
                add_7to5_options_block(doc, opts)

        if method and answer_records is not None:
            group_answers = [
                rec for rec in answer_records
                if rec["exam_id"] == exam_id and rec.get("source_file", "") == source_file
            ]
            add_answer_block(doc, group_answers)
        else:
            add_answer_block(doc, records)

    return doc


def build_reading_doc(groups: dict, subtype: str | None = None) -> Document:
    """
    整篇模式（无 subtype）：
      passage_text → 空行 → 每题题干/选项 → 答案汇总块

    细分模式（有 subtype）：
      每题题干/选项 → 题间空行 → 答案汇总块
    """
    doc = make_document()
    first_unit = True
    global_subtypes = {"main_idea_global", "main_idea_title"}

    for (exam_id, source_file), records in sorted(groups.items()):
        if subtype:
            records = [r for r in records if r.get("question_subtype") == subtype]
        if not records:
            continue

        if not first_unit:
            add_page_break(doc)
        first_unit = False

        if subtype:
            for idx, rec in enumerate(records):
                if subtype in global_subtypes:
                    context = rec.get("passage_text", "")
                    if context:
                        paras = [p for p in context.split("\n") if p.strip()]
                        for p in paras:
                            add_passage_para(doc, p, after=SP_AFTER)
                    else:
                        add_passage_para(doc, f"[{exam_id} 文章原文缺失]", after=SP_AFTER)
                else:
                    excerpt = (rec.get("source_span") or {}).get("excerpt", "")
                    if excerpt:
                        add_passage_para(doc, excerpt, after=SP_AFTER)

                spacer = doc.add_paragraph()
                set_spacing(spacer, before=0, after=SP_AFTER_PASSAGE)
                add_stem_para(doc, rec["source_question_number"], rec.get("question_text", ""))
                add_options_para(doc, rec.get("options") or [], has_answer=False)
                if idx != len(records) - 1:
                    spacer = doc.add_paragraph()
                    set_spacing(spacer, before=0, after=SP_AFTER_PASSAGE)
        else:
            passage = records[0].get("passage_text", "")
            if passage:
                paras = [p for p in passage.split("\n") if p.strip()]
                for p in paras:
                    add_passage_para(doc, p, after=SP_AFTER)
            else:
                add_passage_para(doc, f"[{exam_id} 文章原文缺失]", after=SP_AFTER)

            spacer = doc.add_paragraph()
            set_spacing(spacer, before=0, after=SP_AFTER_PASSAGE)

            for rec in records:
                add_stem_para(doc, rec["source_question_number"], rec.get("question_text", ""))
                add_options_para(doc, rec.get("options") or [], has_answer=False)

        add_answer_block(doc, records)

    return doc


def main():
    parser = argparse.ArgumentParser(description="EXAM_BANKFLOW 导出脚本")
    parser.add_argument("--type",     required=True, choices=["cloze", "reading", "reading_7to5"])
    parser.add_argument("--exam",     help="指定 exam_id（不填则全量）")
    parser.add_argument("--out",      required=True, help="输出 .docx 路径")
    parser.add_argument("--subtype",  help="按 question_subtype 过滤（细分模式）")
    parser.add_argument("--strategy", help="按 strategy 过滤（完形填空细分模式）")
    parser.add_argument("--method",   help="按 solution_methods 过滤（七选五细分模式）")
    parser.add_argument("--limit",    type=int, help="最多输出 N 道题")
    parser.add_argument("--include-no-answer", action="store_true",
                        help="包含 answer=null 的记录（默认只输出有答案的）")
    args = parser.parse_args()

    records = load_records(JSONL_PATH)

    # 筛选题型
    records = [r for r in records if r.get("section_type") == args.type]
    if args.exam:
        records = [r for r in records if r.get("exam_id") == args.exam]

    # 题型专属脏数据过滤
    before = len(records)
    if args.type == "reading_7to5":
        records = [r for r in records if r.get("question_text") and r.get("options")]
    elif args.type == "reading":
        records = [r for r in records if r.get("question_text") and r.get("options")]
    else:
        records = [r for r in records if r.get("question_text")]
    skipped = before - len(records)
    if skipped:
        print(f"跳过脏数据：{skipped} 条")

    if not args.include_no_answer:
        before_ans = len(records)
        records = [r for r in records if r.get("answer")]
        skipped_ans = before_ans - len(records)
        if skipped_ans:
            print(f"跳过无答案记录：{skipped_ans} 条（用 --include-no-answer 关闭过滤）")

    if args.subtype:
        records = [r for r in records if r.get("question_subtype") == args.subtype]
    if args.strategy:
        records = [r for r in records if r.get("strategy") == args.strategy]

    method_answer_records = None

    if args.type == "reading_7to5" and args.method:
        all_groups: dict[tuple, list] = defaultdict(list)
        for r in records:
            key = (r["exam_id"], r.get("source_file", ""))
            all_groups[key].append(r)
        for key in all_groups:
            all_groups[key].sort(key=lambda r: parse_qnum(r.get("source_question_number")))
        groups, method_answer_records = prepare_7to5_method_groups(all_groups, args.method, args.limit)
        records = [rec for group in groups.values() for rec in group]
    else:
        if args.method:
            records = [r for r in records if args.method in (r.get("solution_methods") or [])]
        if args.limit:
            records = records[:args.limit]

        groups: dict[tuple, list] = defaultdict(list)
        for r in records:
            key = (r["exam_id"], r.get("source_file", ""))
            groups[key].append(r)
        for key in groups:
            groups[key].sort(key=lambda r: parse_qnum(r.get("source_question_number")))

    if not records:
        print(f"没有匹配的记录（type={args.type}, exam={args.exam}）", file=sys.stderr)
        sys.exit(1)

    answer_count = len(method_answer_records) if method_answer_records is not None else len(records)
    print(f"共 {answer_count} 条目标题目，扩展后 {len(records)} 条记录，{len(groups)} 个文章单元", flush=True)

    if args.type == "cloze":
        doc = build_cloze_doc(groups, args.subtype, args.strategy)
    elif args.type == "reading":
        doc = build_reading_doc(groups, subtype=args.subtype)
    else:
        doc = build_7to5_doc(groups, method=args.method, answer_records=method_answer_records)

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"✓ 已输出：{out}")


if __name__ == "__main__":
    main()
