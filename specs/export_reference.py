#!/usr/bin/env python3
# =====================================
# !! FORMAT LOCK — 格式基准版本 !!
# 本脚本是唯一授权的Word导出实现
# 任何格式修改必须改此文件并同步specs/export_reference.py
# 禁止agent在调用时绕过此脚本自行生成docx
# 上次验证：E005 完形填空，4项格式全部通过
# =====================================
"""
scripts/export.py — 统一导出入口
用法：
  python scripts/export.py --type cloze --exam E005 --out /tmp/out.docx
  python scripts/export.py --type cloze --out /tmp/out.docx
  python scripts/export.py --type reading --exam E005 --out /tmp/out.docx
  python scripts/export.py --type reading_7to5 --exam E005 --out /tmp/out.docx
  python scripts/export.py --type cloze --subtype cloze_lexical --out /tmp/out.docx
  python scripts/export.py --type cloze --strategy context_clue --out /tmp/out.docx
"""

import argparse
import json
import re
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Twips
from docx.enum.section import WD_ORIENT

# ── 硬编码格式常量 ────────────────────────────────────────────────────────────
FONT_NAME   = "Times New Roman"
FONT_SIZE   = Pt(12)
PAGE_W      = 11906   # A4 宽（DXA）
PAGE_H      = 16838   # A4 高（DXA）
MARGIN      = 1440    # 四边边距（DXA）
LINE_RULE   = "auto"
LINE_VAL    = 360     # 1.5倍行距
SP_BEFORE   = 0
SP_AFTER    = 0
SP_AFTER_Q  = 200     # 答案行后间距
SP_AFTER_PASSAGE = 80   # 文章与题目块之间

TAB_STOPS   = [2200, 4400, 6600]  # 选项列位置（DXA），B/C/D 跳转点

JSONL_PATH  = Path(__file__).resolve().parent.parent / "workflow/records/EXAM-CLEAN-010_STRUCTURED_QUESTION_BANK.jsonl"
QNUM_RE = re.compile(r"\d+")


# ── 工具函数 ──────────────────────────────────────────────────────────────────

def set_spacing(para, before=SP_BEFORE, after=SP_AFTER, line=LINE_VAL):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"),   str(before))
    spacing.set(qn("w:after"),    str(after))
    spacing.set(qn("w:line"),     str(line))
    spacing.set(qn("w:lineRule"), LINE_RULE)
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)


def add_tab_stops(para):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    for pos in TAB_STOPS:
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), str(pos))
        tabs.append(tab)
    existing = pPr.find(qn("w:tabs"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(tabs)


def make_run(para, text, bold=False):
    run = para.add_run(text)
    run.font.name       = FONT_NAME
    run.font.size       = FONT_SIZE
    run.font.bold       = bold
    # 中文字体
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    FONT_NAME)
    rFonts.set(qn("w:hAnsi"),    FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    return run


# 完形填空占位符：支持 ____21____ / ___21___ / 空格21空格 三种格式
_CLOZE_BLANK_RE = re.compile(r'(_{2,}|\s{2,})(\d{2})(_{2,}|\s{2,})')


def add_passage_para(doc, text, after=SP_AFTER, underline_word: str | None = None,
                     cloze_blanks: bool = False):
    """添加文章段落。
    underline_word: 词汇题下划线词。
    cloze_blanks:   完形填空模式，将 '   41   ' 渲染为带下划线的空白+题号。
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, before=SP_BEFORE, after=after)
    full_text = "　　" + text  # 首行缩进两个全角空格

    if cloze_blanks and _CLOZE_BLANK_RE.search(full_text):
        # 按完形空格切分，加下划线
        pos = 0
        for m in _CLOZE_BLANK_RE.finditer(full_text):
            if m.start() > pos:
                make_run(para, full_text[pos:m.start()])
            # 渲染为 "______41______"（下划线 run）
            blank_text = "______" + m.group(2) + "______"
            r = make_run(para, blank_text)
            r.underline = True
            pos = m.end()
        if pos < len(full_text):
            make_run(para, full_text[pos:])
    elif underline_word and underline_word.lower() in full_text.lower():
        idx = full_text.lower().index(underline_word.lower())
        before_part = full_text[:idx]
        word_part   = full_text[idx: idx + len(underline_word)]
        after_part  = full_text[idx + len(underline_word):]
        if before_part:
            make_run(para, before_part)
        r = make_run(para, word_part)
        r.underline = True
        if after_part:
            make_run(para, after_part)
    else:
        make_run(para, full_text)
    return para


def add_stem_para(doc, qnum, text, tag=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, before=120, after=0)
    qnum_str = str(qnum).rstrip(".")
    tag_str = f"【{tag}】" if tag else ""
    make_run(para, f"{qnum_str}.{tag_str} {text}")
    return para


def add_options_para(doc, options, has_answer, qnum=None):
    """完形填空用：ABCD tab分列，一行输出（选项较短）。"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_tab_stops(para)
    after = 0 if has_answer else SP_AFTER_Q
    set_spacing(para, before=120, after=after)

    labels = ["A", "B", "C", "D"]
    opt_map = {o["label"]: o["text"] for o in options}
    prefix = f"{qnum}. " if qnum is not None else ""
    text = prefix + "\t".join(f"{lbl}. {opt_map.get(lbl, '')}" for lbl in labels)
    make_run(para, text)
    return para


def add_reading_options_paras(doc, options, incomplete=False):
    """阅读理解用：ABCD 每项单独一行（选项为完整句子）。
    incomplete=True 时末尾加警告行。"""
    labels = ["A", "B", "C", "D"]
    opt_map = {o["label"]: o["text"] for o in options}
    for lbl in labels:
        text = opt_map.get(lbl)
        if text is None:
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(para, before=0, after=0)
        make_run(para, f"{lbl}. {text}")
    if incomplete:
        warn = doc.add_paragraph()
        warn.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(warn, before=0, after=0)
        r = warn.add_run("[⚠ 选项不完整，需补充]")
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE
        r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def add_answer_para(doc, answer):
    ans_text = answer.strip() if (answer and answer.strip()) else "待补充"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(para, before=0, after=SP_AFTER_Q)
    make_run(para, f"【答案】{ans_text}", bold=True)
    return para


def add_answer_block(doc, records, per_line=False):
    """
    答案汇总块。
    per_line=False（完形/整篇）：每5题一行，如 "21-25  A B C D A"
    per_line=True （阅读细分）：每题单独一行，如 "4. D"
    """
    sorted_recs = sorted(records, key=lambda r: parse_qnum(r.get("source_question_number")))

    spacer = doc.add_paragraph()
    set_spacing(spacer, before=0, after=0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(title, before=0, after=0)
    make_run(title, "【答案】", bold=True)

    if per_line:
        for r in sorted_recs:
            num = str(r.get("source_question_number", "")).rstrip(".")
            ans = (r.get("answer") or "").strip() or "_"
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_spacing(para, before=0, after=0)
            make_run(para, f"{num}. {ans}", bold=True)
    else:
        for i in range(0, len(sorted_recs), 5):
            chunk = sorted_recs[i:i + 5]
            nums = [parse_qnum(r.get("source_question_number")) for r in chunk]
            answers = [(r.get("answer") or "").strip() or "_" for r in chunk]
            label = str(nums[0]) if len(nums) == 1 else f"{nums[0]}-{nums[-1]}"
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_spacing(para, before=0, after=0)
            make_run(para, f"{label}  " + " ".join(answers), bold=True)


def add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)
    set_spacing(para, before=0, after=0)


# ── 文档初始化 ────────────────────────────────────────────────────────────────

def make_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Twips(PAGE_W)
    sec.page_height   = Twips(PAGE_H)
    sec.top_margin    = Twips(MARGIN)
    sec.bottom_margin = Twips(MARGIN)
    sec.left_margin   = Twips(MARGIN)
    sec.right_margin  = Twips(MARGIN)
    # 清除默认样式间距
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.space_before = 0
    style.paragraph_format.space_after  = 0
    return doc


# ── 主逻辑 ────────────────────────────────────────────────────────────────────

def load_records(path: Path) -> list:
    with open(path, encoding="utf-8") as f:
        return [json.loads(l) for l in f if l.strip()]


def parse_qnum(value) -> int:
    match = QNUM_RE.search(str(value or ""))
    return int(match.group()) if match else 0


def build_cloze_doc(groups: dict, subtype: str | None, strategy: str | None) -> Document:
    """
    groups: { (exam_id, source_file): [records sorted by q_number] }
    """
    doc        = make_document()
    first_unit = True

    for (exam_id, source_file), records in sorted(groups.items()):
        # 细分过滤
        if subtype:
            records = [r for r in records if r.get("question_subtype") == subtype]
        if strategy:
            records = [r for r in records if r.get("strategy") == strategy]
        if not records:
            continue

        if not first_unit:
            if subtype or strategy:
                sep = doc.add_paragraph()
                set_spacing(sep, before=0, after=0, line=LINE_VAL)
            else:
                add_page_break(doc)
        first_unit = False

        use_excerpt = bool(subtype or strategy)

        if use_excerpt:
            # 细分模式：excerpt 上下文 + 选项行（含序号），末尾汇总答案
            for rec in records:
                excerpt = rec.get("source_span", {}).get("excerpt", rec.get("question_text", ""))
                add_passage_para(doc, excerpt, after=SP_AFTER_PASSAGE, cloze_blanks=True)
                add_options_para(doc, rec.get("options") or [], has_answer=False,
                                 qnum=rec["source_question_number"])
            add_answer_block(doc, records)
        else:
            # 整篇模式：文章原文 → 所有题目
            passage = records[0].get("passage_text", "")
            if passage:
                paras = [p for p in passage.split("\n") if p.strip()]
                for i, p in enumerate(paras):
                    after = SP_AFTER_PASSAGE if i == len(paras) - 1 else SP_AFTER
                    add_passage_para(doc, p, after=after, cloze_blanks=True)
            else:
                add_passage_para(doc, f"[{exam_id} 文章原文缺失]", after=SP_AFTER_PASSAGE)

            for rec in records:
                add_options_para(doc, rec.get("options") or [], has_answer=False,
                                 qnum=rec["source_question_number"])
            add_answer_block(doc, records)

    return doc


def add_7to5_options_block(doc, options):
    """七选五选项池：A-G 每项单独一行（句子较长，不分列）。"""
    spacer = doc.add_paragraph()
    set_spacing(spacer, before=0, after=0)
    for opt in options:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(para, before=0, after=0)
        make_run(para, f"{opt['label']}. {opt['text']}")


def split_passage_paragraphs(passage: str) -> list[str]:
    return [p.strip() for p in (passage or "").split("\n") if p.strip()]


def get_7to5_placeholder(rec: dict) -> str | None:
    """返回该题在 passage_text 中的占位符，支持 ___N___ 和 ____N____ 两种格式。"""
    question_text = rec.get("question_text") or ""
    match = re.search(r"_{3,}(\d+)_{3,}", question_text)
    if match:
        return match.group(0)  # 保留原始格式
    qnum = parse_qnum(rec.get("source_question_number"))
    return str(qnum) if qnum else None  # 退化：直接用题号字符串匹配


def find_7to5_para_index(paras: list[str], rec: dict) -> int:
    qnum = parse_qnum(rec.get("source_question_number"))
    if qnum:
        # 支持 ___N___ / ____N____ 两种格式（避免 f-string 吃掉 {3,}）
        pat = re.compile(r"_{3,}" + str(qnum) + r"_{3,}")
        for idx, para in enumerate(paras):
            if pat.search(para):
                return idx
    return max(len(paras) - 1, 0)


def add_7to5_marker_line(doc, records: list[dict]):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(para, before=0, after=SP_AFTER)
    chunks = []
    for rec in sorted(records, key=lambda r: parse_qnum(r.get("source_question_number"))):
        qnum = parse_qnum(rec.get("source_question_number"))
        methods = rec.get("solution_methods") or []
        label = "·".join(methods) if methods else "unlabeled"
        chunks.append(f"{qnum}【{label}】")
    make_run(para, "▶ " + "  ".join(chunks), bold=True)


def prepare_7to5_method_groups(groups: dict, method: str, limit: int | None):
    selected_ids: list[str] = []
    selected_set: set[str] = set()
    expanded_groups: dict[tuple, list] = {}

    for key in sorted(groups):
        group_records = groups[key]
        paras = split_passage_paragraphs(group_records[0].get("passage_text", ""))
        para_map: dict[int, list[dict]] = defaultdict(list)
        for rec in group_records:
            para_index = find_7to5_para_index(paras, rec)
            rec["_export_para_index"] = para_index
            para_map[para_index].append(rec)

        matched = [
            rec for rec in group_records
            if method in (rec.get("solution_methods") or [])
        ]
        matched.sort(key=lambda r: parse_qnum(r.get("source_question_number")))

        included_ids: set[str] = set()
        for rec in matched:
            rec_id = rec.get("question_id") or f"{key}-{rec.get('source_question_number')}"
            if limit is not None and len(selected_ids) >= limit:
                break
            selected_ids.append(rec_id)
            selected_set.add(rec_id)
            para_index = rec["_export_para_index"]
            for block_rec in para_map.get(para_index, []):
                block_id = block_rec.get("question_id") or f"{key}-{block_rec.get('source_question_number')}"
                included_ids.add(block_id)

        expanded = [
            rec for rec in group_records
            if (rec.get("question_id") or f"{key}-{rec.get('source_question_number')}") in included_ids
        ]
        expanded.sort(key=lambda r: (r["_export_para_index"], parse_qnum(r.get("source_question_number"))))
        if expanded:
            expanded_groups[key] = expanded

        if limit is not None and len(selected_ids) >= limit:
            # limit applies to method-matched questions; later groups are ignored
            break

    answer_records = []
    for key in sorted(expanded_groups):
        for rec in expanded_groups[key]:
            rec_id = rec.get("question_id") or f"{key}-{rec.get('source_question_number')}"
            if rec_id in selected_set:
                answer_records.append(rec)

    return expanded_groups, answer_records


def build_7to5_doc(groups: dict, method: str | None = None, answer_records: list[dict] | None = None) -> Document:
    """
    整篇模式（无 method 筛选）：
      passage_text → A-G 选项池 → 答案汇总块

    细分模式（有 method 筛选）：
      按段落块输出：从首段到空格所在段 → A-G 选项池 → 标注行；末尾统一答案汇总块
    """
    doc        = make_document()
    first_unit = True

    for (exam_id, source_file), records in sorted(groups.items()):
        if not records:
            continue
        if not first_unit:
            if method:
                sep = doc.add_paragraph()
                set_spacing(sep, before=0, after=0, line=LINE_VAL)
            else:
                add_page_break(doc)
        first_unit = False

        opts = records[0].get("options") or []

        if method:
            paras = split_passage_paragraphs(records[0].get("passage_text", ""))
            blocks: dict[int, list[dict]] = defaultdict(list)
            for rec in records:
                para_index = rec.get("_export_para_index")
                if para_index is None:
                    para_index = find_7to5_para_index(paras, rec)
                blocks[para_index].append(rec)

            sorted_para_indices = sorted(blocks)
            # 文章中第一个含空白的段落（扫描全篇 passage，与导出集合无关）
            first_blank_para = next(
                (i for i, p in enumerate(paras) if re.search(r"____\d+____", p)),
                0
            )

            for block_idx, para_index in enumerate(sorted_para_indices):
                block_records = sorted(
                    blocks[para_index],
                    key=lambda r: parse_qnum(r.get("source_question_number"))
                )
                if paras:
                    # 该题所在段是文章第一个空白段：输出 paras[0..N]（需要前文上下文）
                    # 否则：只输出 paras[N]（该段即可）
                    if para_index == first_blank_para:
                        for idx in range(para_index + 1):
                            after = SP_AFTER if idx < para_index else 0
                            add_passage_para(doc, paras[idx], after=after)
                    else:
                        add_passage_para(doc, paras[para_index], after=0)
                else:
                    add_passage_para(doc, f"[{exam_id} 文章原文缺失]", after=0)

                spacer = doc.add_paragraph()
                set_spacing(spacer, before=0, after=0)
                if opts:
                    add_7to5_options_block(doc, opts)

                if block_idx != len(blocks) - 1:
                    spacer = doc.add_paragraph()
                    set_spacing(spacer, before=0, after=SP_AFTER_PASSAGE)
        else:
            passage = records[0].get("passage_text", "")
            if passage:
                paras = [p for p in passage.split("\n") if p.strip()]
                for i, p in enumerate(paras):
                    after = SP_AFTER_PASSAGE if i == len(paras) - 1 else SP_AFTER
                    add_passage_para(doc, p, after=after)
            else:
                add_passage_para(doc, f"[{exam_id} 文章原文缺失]", after=SP_AFTER_PASSAGE)

            if opts:
                add_7to5_options_block(doc, opts)

        if method and answer_records is not None:
            group_answers = [
                rec for rec in answer_records
                if rec["exam_id"] == exam_id and rec.get("source_file", "") == source_file
            ]
            add_answer_block(doc, group_answers)
        else:
            add_answer_block(doc, records)

    return doc


_PARA_REF_RE = re.compile(
    r"paragraph\s+(\d+)"
    r"|the\s+(first|second|third|fourth|fifth|sixth|seventh|eighth|ninth|tenth)\s+paragraph"
    r"|第\s*([一二三四五六七八九十\d]+)\s*段",
    re.IGNORECASE,
)
# 多段：paragraphs N and M / paragraphs N-M / paragraphs N to M
_PARA_RANGE_RE = re.compile(
    r"paragraphs?\s+(\d+)\s+(?:and|to|-)\s+(\d+)",
    re.IGNORECASE,
)
_LAST_PARA_RE = re.compile(
    r"last\s+paragraph|final\s+paragraph|最后[一]?段|末段",
    re.IGNORECASE,
)
# 最后N段：the last two/three/2/3 paragraphs
_LAST_N_PARA_RE = re.compile(
    r"last\s+(two|three|four|2|3|4)\s+paragraphs",
    re.IGNORECASE,
)
_LAST_N_MAP = {"two": 2, "three": 3, "four": 4, "2": 2, "3": 3, "4": 4}
_CN_NUM = {"一":1,"二":2,"三":3,"四":4,"五":5,"六":6,"七":7,"八":8,"九":9,"十":10}
_EN_ORD = {"first":1,"second":2,"third":3,"fourth":4,"fifth":5,
           "sixth":6,"seventh":7,"eighth":8,"ninth":9,"tenth":10}

def _para_ref_from_text(question_text: str) -> list[int] | None:
    """从题目文字中解析明确的段落序号（1-based列表），无法解析返回 None。
    支持单段（paragraph 4）和多段范围（paragraphs 4 and 5 / paragraphs 4-6）。"""
    qt = question_text or ""
    # 先检测多段范围
    mr = _PARA_RANGE_RE.search(qt)
    if mr:
        start, end = int(mr.group(1)), int(mr.group(2))
        return list(range(min(start, end), max(start, end) + 1))
    # 单段
    m = _PARA_REF_RE.search(qt)
    if not m:
        return None
    if m.group(1):
        return [int(m.group(1))]
    if m.group(2):
        n = _EN_ORD.get(m.group(2).lower())
        return [n] if n else None
    if m.group(3):
        s = m.group(3).strip()
        n = int(s) if s.isdigit() else _CN_NUM.get(s)
        return [n] if n else None
    return None


_UNDERLINE_TARGET_RE = re.compile(
    r'(?:underlined\s+(?:word|phrase|expression)\s*["“「]([^"”」]+)["”」]'
    r'|the\s+(?:word|phrase|expression)\s*["“「]([^"”」]+)["”」])',
    re.IGNORECASE,
)


def _extract_underline_target(question_text: str) -> str | None:
    """从题目文字中提取带下划线的目标词/短语，无法解析返回 None。"""
    m = _UNDERLINE_TARGET_RE.search(question_text or "")
    if not m:
        return None
    return (m.group(1) or m.group(2) or "").strip() or None


def _extract_context_paras(passage_text: str, para_index: int | None,
                           window: int = 2) -> list[str]:
    """
    从 passage_text 按段落索引截取上下文（前后 window 段）。
    para_index 缺失时返回空列表（调用方降级处理）。
    """
    if not passage_text or para_index is None:
        return []
    try:
        para_index = int(para_index)
    except (ValueError, TypeError):
        return []
    all_paras = [p for p in passage_text.split("\n") if p.strip()]
    start = max(0, para_index - window)
    end   = min(len(all_paras), para_index + window + 1)
    return all_paras[start:end]


def build_reading_doc(groups: dict, subtype: str | None = None) -> Document:
    """
    整篇模式（无 subtype）：
      passage_text → 空行 → 每题题干/选项（每项单独一行） → 答案汇总块

    细分模式（有 subtype）：
      全篇subtype（main_idea_global/title）：完整 passage_text → 题干/选项 → 答案块
      其他subtype：每题从 passage_text 按 para_index 截取前后2段上下文
                   → 题干 → 选项（每项单独一行） → 答案汇总块
    """
    doc = make_document()
    first_unit = True
    global_subtypes = {"main_idea_global", "main_idea_title"}

    for (exam_id, source_file), records in sorted(groups.items()):
        if subtype:
            records = [r for r in records
                       if (r.get("question_subtype") or "").startswith(subtype)]
        if not records:
            continue

        if not first_unit:
            if subtype:
                # 细分模式：exam单元之间空一行，不分页
                sep = doc.add_paragraph()
                set_spacing(sep, before=0, after=0, line=LINE_VAL)
            else:
                add_page_break(doc)
        first_unit = False

        # 整篇模式用 records[0] 的文章；细分模式每题用自己的 passage_text
        passage = records[0].get("passage_text", "")
        passage_paras = [p for p in passage.split("\n") if p.strip()] if passage else []

        if subtype:
            # 细分模式：每题用自己的 passage_text
            for idx_rec, rec in enumerate(records):
                # 题与题之间空一行
                if idx_rec > 0:
                    sep = doc.add_paragraph()
                    set_spacing(sep, before=0, after=0, line=LINE_VAL)

                opts = rec.get("options") or []
                incomplete = len(opts) < 4
                qt = rec.get("question_text", "")
                ref = _para_ref_from_text(qt)
                is_last = bool(_LAST_PARA_RE.search(qt or ""))
                last_n_m = _LAST_N_PARA_RE.search(qt or "")
                last_n = _LAST_N_MAP.get(last_n_m.group(1).lower()) if last_n_m else None
                underline_word = _extract_underline_target(qt)

                rec_passage = rec.get("passage_text", "")
                rec_paras = [p for p in rec_passage.split("\n") if p.strip()] if rec_passage else []

                def _add_paras(paras):
                    for p in paras:
                        uw = underline_word if (underline_word and underline_word.lower() in p.lower()) else None
                        add_passage_para(doc, p, after=SP_AFTER, underline_word=uw)

                if last_n and rec_paras:
                    _add_paras(rec_paras[-last_n:])
                elif is_last and rec_paras:
                    _add_paras([rec_paras[-1]])
                elif ref is not None and rec_paras:
                    # ref 是1-based段落号列表（单段或多段）
                    selected = [rec_paras[i - 1] for i in ref if 0 < i <= len(rec_paras)]
                    _add_paras(selected if selected else rec_paras)  # 越界时降级为全文
                elif rec_paras:
                    _add_paras(rec_paras)
                else:
                    add_passage_para(doc, f"[{exam_id} 文章原文缺失]", after=SP_AFTER)

                spacer = doc.add_paragraph()
                set_spacing(spacer, before=0, after=SP_AFTER_PASSAGE)
                add_stem_para(doc, rec["source_question_number"], qt)
                add_reading_options_paras(doc, opts, incomplete=incomplete)
        else:
            # 整篇模式：全文出一次，所有题目跟在后面
            if passage_paras:
                for p in passage_paras:
                    add_passage_para(doc, p, after=SP_AFTER)
            else:
                add_passage_para(doc, f"[{exam_id} 文章原文缺失]", after=SP_AFTER)

            spacer = doc.add_paragraph()
            set_spacing(spacer, before=0, after=SP_AFTER_PASSAGE)

            for rec in records:
                opts = rec.get("options") or []
                add_stem_para(doc, rec["source_question_number"], rec.get("question_text", ""))
                add_reading_options_paras(doc, opts, incomplete=len(opts) < 4)

        add_answer_block(doc, records, per_line=bool(subtype))

    return doc


def main():
    parser = argparse.ArgumentParser(description="EXAM_BANKFLOW 导出脚本")
    parser.add_argument("--type",     required=True, choices=["cloze", "reading", "reading_7to5"])
    parser.add_argument("--exam",     help="指定 exam_id（不填则全量）")
    parser.add_argument("--out",      required=True, help="输出 .docx 路径")
    parser.add_argument("--subtype",  help="按 question_subtype 过滤（细分模式）")
    parser.add_argument("--strategy", help="按 strategy 过滤（完形填空细分模式）")
    parser.add_argument("--method",   help="按 solution_methods 过滤（七选五细分模式）")
    parser.add_argument("--limit",    type=int, help="最多输出 N 道题")
    parser.add_argument("--include-no-answer", action="store_true",
                        help="包含 answer=null 的记录（默认只输出有答案的）")
    args = parser.parse_args()

    records = load_records(JSONL_PATH)

    # 筛选题型
    records = [r for r in records if r.get("section_type") == args.type]
    if args.exam:
        records = [r for r in records if r.get("exam_id") == args.exam]

    # 题型专属脏数据过滤
    before = len(records)
    if args.type == "reading_7to5":
        records = [r for r in records if r.get("question_text") and r.get("options")]
    elif args.type == "reading":
        records = [r for r in records if r.get("question_text") and r.get("options")]
    else:
        records = [r for r in records if r.get("question_text")]
    skipped = before - len(records)
    if skipped:
        print(f"跳过脏数据：{skipped} 条")

    if not args.include_no_answer:
        before_ans = len(records)
        records = [r for r in records if r.get("answer")]
        skipped_ans = before_ans - len(records)
        if skipped_ans:
            print(f"跳过无答案记录：{skipped_ans} 条（用 --include-no-answer 关闭过滤）")

    if args.subtype:
        # 前缀匹配：--subtype inference 匹配 inference_contrast / inference_logic 等
        records = [r for r in records
                   if (r.get("question_subtype") or "").startswith(args.subtype)]
    if args.strategy:
        records = [r for r in records if r.get("strategy") == args.strategy]

    method_answer_records = None

    if args.type == "reading_7to5" and args.method:
        all_groups: dict[tuple, list] = defaultdict(list)
        for r in records:
            key = (r["exam_id"], r.get("source_file", ""))
            all_groups[key].append(r)
        for key in all_groups:
            all_groups[key].sort(key=lambda r: parse_qnum(r.get("source_question_number")))
        groups, method_answer_records = prepare_7to5_method_groups(all_groups, args.method, args.limit)
        records = [rec for group in groups.values() for rec in group]
    else:
        if args.method:
            records = [r for r in records if args.method in (r.get("solution_methods") or [])]
        if args.limit:
            records = records[:args.limit]

        groups: dict[tuple, list] = defaultdict(list)
        for r in records:
            key = (r["exam_id"], r.get("source_file", ""))
            groups[key].append(r)
        for key in groups:
            groups[key].sort(key=lambda r: parse_qnum(r.get("source_question_number")))

    if not records:
        print(f"没有匹配的记录（type={args.type}, exam={args.exam}）", file=sys.stderr)
        sys.exit(1)

    answer_count = len(method_answer_records) if method_answer_records is not None else len(records)
    print(f"共 {answer_count} 条目标题目，扩展后 {len(records)} 条记录，{len(groups)} 个文章单元", flush=True)

    if args.type == "cloze":
        doc = build_cloze_doc(groups, args.subtype, args.strategy)
    elif args.type == "reading":
        doc = build_reading_doc(groups, subtype=args.subtype)
    else:
        doc = build_7to5_doc(groups, method=args.method, answer_records=method_answer_records)

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"✓ 已输出：{out}")


if __name__ == "__main__":
    main()
